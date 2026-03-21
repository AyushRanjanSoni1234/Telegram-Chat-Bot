import os
import json
import numpy as np
import pandas as pd
from dotenv import load_dotenv
from huggingface_hub import InferenceClient

from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# PDF support
from PyPDF2 import PdfReader

# ✅ FIX: Added docx support
from docx import Document as DocxDocument  

load_dotenv()
HF_TOKEN = os.getenv("API_KEY")


# -------------------------------
# 🔥 Custom Embedding Class
# -------------------------------
class HFEmbedding:
    def __init__(self, token):
        self.client = InferenceClient(token=token)
        self.model = "sentence-transformers/all-MiniLM-L6-v2"

    def embed_query(self, text):
        try:
            emb = self.client.feature_extraction(text, model=self.model)

            # ✅ FIX: Handle empty or invalid response
            if emb is None or len(emb) == 0:
                return [0.0] * 384

            # ✅ FIX: Handle nested structure safely
            vector = np.array(emb, dtype="float32")

            # ✅ FIX: Ensure vector is valid
            if vector is None or len(vector) == 0:
                return [0.0] * 384

            return vector.tolist()

        except Exception as e:
            print(f"❌ Embedding error: {e}")
            return [0.0] * 384  # fallback

    def embed_documents(self, texts):
        # ✅ FIX: Filter empty texts
        return [self.embed_query(t) for t in texts if t.strip()]


# -------------------------------
# ✂️ Chunk Function
# -------------------------------
def chunk_text(text, size=300):
    return [text[i:i+size] for i in range(0, len(text), size) if text[i:i+size].strip()]  # ✅ FIX


# -------------------------------
# 📂 Load Documents (All Types)
# -------------------------------
def load_docs(folder="Data"):
    docs = []

    # ✅ FIX: Check folder exists
    if not os.path.exists(folder):
        raise ValueError(f"❌ Folder '{folder}' not found")

    for file in os.listdir(folder):
        file_path = os.path.join(folder, file)

        try:
            # TXT
            if file.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    docs.append((file, f.read()))

            # ✅ FIX: Proper DOCX handling (NOT open())
            elif file.endswith((".docx", ".doc")):
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
                docs.append((file, text))

            # PDF
            elif file.endswith(".pdf"):
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                docs.append((file, text))

            # CSV
            elif file.endswith(".csv"):
                df = pd.read_csv(file_path)
                docs.append((file, df.to_string()))

            # JSON
            elif file.endswith(".json"):
                with open(file_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    docs.append((file, json.dumps(data)))

        except Exception as e:
            print(f"❌ Error loading {file}: {e}")

    return docs


# -------------------------------
# 📄 Create Documents
# -------------------------------
def create_documents():
    raw_docs = load_docs()
    documents = []

    for filename, doc in raw_docs:

        # ✅ FIX: Skip empty docs
        if not doc or not doc.strip():
            continue

        chunks = chunk_text(doc, size=300)

        for chunk in chunks:
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={"source": filename}
                )
            )

    return documents


# -------------------------------
# 🧠 Create FAISS Index
# -------------------------------
def create_index():
    documents = create_documents()

    # ✅ FIX: Check before FAISS
    if len(documents) == 0:
        raise ValueError("❌ No documents found. Check your Data folder.")

    print(f"✅ Total chunks created: {len(documents)}")  # DEBUG

    embeddings = HFEmbedding(token=HF_TOKEN)

    db = FAISS.from_documents(documents, embeddings)

    db.save_local("db")

    print("✅ Index Created Successfully (HuggingFace Embeddings)")


# -------------------------------
# 🚀 Run
# -------------------------------
if __name__ == "__main__":
    create_index()